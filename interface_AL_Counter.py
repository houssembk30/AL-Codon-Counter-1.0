# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'AL-Counter.ui'
#
# Created by: PyQt5 UI code generator 5.13.0
# Houssem Ben Kalfallah 2024
#
# WARNING! All changes made in this file will be lost!


from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog, QMainWindow, QApplication, QPushButton, QLabel, QMessageBox
from AL_CODON_test_2 import AL_CODON_Counter
from Read_Docx_and_generate_Txt__Docx import Read_Doc_and_generate_Txt
from View_Pentamer_Colored_Text import Ui_MainWindow_2
from View_RED_or_Colored_Selected_Sequence import Ui_MainWindow_3
from splash_screen import SplashScreen

from docx import Document
from docx.shared import RGBColor, Pt
from bs4 import BeautifulSoup
import matplotlib.colors as mcolors
import re
import sys
# interacting with the file system, managing processes, and accessing the environment.
import os

# ---------------------------------------------
# Lib added for Multiprocessing and Thread
from PyQt5.QtCore import QThread, pyqtSignal, QTimer
from multiprocessing import Pool, Manager, freeze_support

import time


# ---------------------------------------------
# ----------> By keeping process_data outside of any class, we ensure it remains simple, easily picklable,
# and compatible with the multiprocessing module's requirements. This separation of concerns also enhances code clarity,
# making it clear that process_data is purely a data processing function, not tied to the state or behavior
# of any class instance.
#
# Function to process data
# This function simulates heavy data processing
# Parameters: Takes a tuple args containing title, sequence_data, and additional_text.
# Returns: A tuple (title, processed_sequence, additional_info)
# where processed_sequence is the processed sequence data and additional_info is a string from additional_text.
def process_data(args):
    title, sequence_data, additional_text = args
    # Simulate heavy processing
    processed_sequence = sequence_data[title]
    additional_info = additional_text
    # Simulate delay
    time.sleep(1)
    return title, processed_sequence, additional_info


# using the resource_path function to ensure the icon path is correctly located in both the development and bundled
# environments.
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


# This thread continuously checks for processed results in a queue and emits a signal when new data is available.
class ResultListenerThread(QThread):
    data_processed = pyqtSignal(str, str, str)

    # Attributes:
    # results_queue: A queue to hold the results from the worker processes.
    # data_processed: A PyQt signal that emits the processed data.
    # _is_running: A flag to control the running state of the thread.

    def __init__(self, results_queue):
        super().__init__()
        self.results_queue = results_queue
        self._is_running = True

    # run: The main loop that runs in the thread,
    # checking the results_queue for new results and emitting the data_processed signal.
    def run(self):
        while self._is_running:
            if not self.results_queue.empty():
                result = self.results_queue.get()
                self.data_processed.emit(*result)
            time.sleep(0.1)

    # stop: Stops the thread by setting _is_running to False.
    def stop(self):
        self._is_running = False


class Ui_AL_Codon_Counter(object):
    def setupUi(self, AL_Codon_Counter):
        AL_Codon_Counter.setObjectName("AL_Codon_Counter")
        AL_Codon_Counter.resize(879, 599)
        # Construct the path to the image
        icon_path = resource_path("JPG_PNG IMAGES/logo.png")
        AL_Codon_Counter.setWindowIcon(QtGui.QIcon(icon_path))

        self.centralwidget = QtWidgets.QWidget(AL_Codon_Counter)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 0, 0, 1, 2)
        self.liste_pentamers_comboBox = QtWidgets.QComboBox(self.centralwidget)
        self.liste_pentamers_comboBox.setObjectName("liste_pentamers_comboBox")
        self.gridLayout.addWidget(self.liste_pentamers_comboBox, 7, 2, 1, 1)
        self.liste_pentamers_comboBox_5 = QtWidgets.QComboBox(self.centralwidget)
        self.liste_pentamers_comboBox_5.setObjectName("liste_pentamers_comboBox_5")
        self.gridLayout.addWidget(self.liste_pentamers_comboBox_5, 7, 3, 1, 1)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 5, 0, 1, 1)
        self.list_Sequence_titles_comboBox = QtWidgets.QComboBox(self.centralwidget)
        self.list_Sequence_titles_comboBox.setObjectName("list_Sequence_titles_comboBox")
        self.gridLayout.addWidget(self.list_Sequence_titles_comboBox, 10, 0, 1, 1)
        self.pushButton_renamefile = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_renamefile.setObjectName("pushButton_renamefile")
        self.gridLayout.addWidget(self.pushButton_renamefile, 15, 0, 1, 1)
        self.pushButton_launch_counter = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_launch_counter.setObjectName("pushButton_launch_counter")
        self.gridLayout.addWidget(self.pushButton_launch_counter, 15, 2, 1, 1)
        self.label_9 = QtWidgets.QLabel(self.centralwidget)
        self.label_9.setObjectName("label_9")
        self.gridLayout.addWidget(self.label_9, 2, 0, 1, 1)
        self.selected_fileshow = QtWidgets.QLabel(self.centralwidget)
        self.selected_fileshow.setText("")
        self.selected_fileshow.setObjectName("selected_fileshow")
        self.gridLayout.addWidget(self.selected_fileshow, 14, 0, 1, 3)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 0, 3, 1, 1)
        self.pushButton_Create_ALL_in_One = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_Create_ALL_in_One.setObjectName("pushButton_Create_ALL_in_One")
        self.gridLayout.addWidget(self.pushButton_Create_ALL_in_One, 12, 0, 1, 1)
        self.pushButton_generate_Docx_seq = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_generate_Docx_seq.setObjectName("pushButton_generate_Docx_seq")
        self.gridLayout.addWidget(self.pushButton_generate_Docx_seq, 17, 0, 1, 1)
        self.label_8 = QtWidgets.QLabel(self.centralwidget)
        self.label_8.setObjectName("label_8")
        self.gridLayout.addWidget(self.label_8, 11, 0, 1, 1)
        self.pushButton_Next_If_Doc_or_multi_Files = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_Next_If_Doc_or_multi_Files.setObjectName("pushButton_Next_If_Doc_or_multi_Files")
        self.gridLayout.addWidget(self.pushButton_Next_If_Doc_or_multi_Files, 7, 0, 1, 1)
        self.pushButton_generate_Txt_seq = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_generate_Txt_seq.setObjectName("pushButton_generate_Txt_seq")
        self.gridLayout.addWidget(self.pushButton_generate_Txt_seq, 16, 0, 1, 1)
        self.DNA_Sequnce_output = QtWidgets.QTextEdit(self.centralwidget)

        self.DNA_Sequnce_output.setObjectName("DNA_Sequnce_output")
        self.gridLayout.addWidget(self.DNA_Sequnce_output, 1, 3, 1, 1)
        self.textEdit_results_TAIL = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_results_TAIL.setObjectName("textEdit_results_TAIL")
        self.gridLayout.addWidget(self.textEdit_results_TAIL, 4, 3, 1, 1)
        self.radio_red_Button = QtWidgets.QRadioButton(self.centralwidget)
        self.radio_red_Button.setObjectName("radio_red_Button")
        self.gridLayout.addWidget(self.radio_red_Button, 8, 2, 1, 1)
        self.colored_radio_Button = QtWidgets.QRadioButton(self.centralwidget)
        self.colored_radio_Button.setObjectName("colored_radio_Button")
        self.gridLayout.addWidget(self.colored_radio_Button, 9, 2, 1, 1)
        self.colored_radio_Button_tail = QtWidgets.QRadioButton(self.centralwidget)
        self.colored_radio_Button_tail.setObjectName("colored_radio_Button_tail")
        self.gridLayout.addWidget(self.colored_radio_Button_tail, 11, 2, 1, 1)
        self.red_radio_Button_tail = QtWidgets.QRadioButton(self.centralwidget)
        self.red_radio_Button_tail.setObjectName("red_radio_Button_tail")
        self.gridLayout.addWidget(self.red_radio_Button_tail, 10, 2, 1, 1)

        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setObjectName("label_7")
        self.gridLayout.addWidget(self.label_7, 8, 0, 1, 1)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 13, 0, 1, 1)
        self.textEdit_results = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_results.setObjectName("textEdit_results")
        self.gridLayout.addWidget(self.textEdit_results, 10, 3, 7, 1)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 8, 3, 1, 1)
        self.label_10 = QtWidgets.QLabel(self.centralwidget)
        self.label_10.setObjectName("label_10")
        self.gridLayout.addWidget(self.label_10, 2, 3, 1, 1)
        self.label_6 = QtWidgets.QLabel(self.centralwidget)
        self.label_6.setObjectName("label_6")
        self.gridLayout.addWidget(self.label_6, 5, 2, 1, 1)
        self.label_5_Pentamers = QtWidgets.QLabel(self.centralwidget)
        self.label_5_Pentamers.setObjectName("label_5_Pentamers")
        self.gridLayout.addWidget(self.label_5_Pentamers, 5, 3, 1, 1)
        self.DNA_Sequnce_input = QtWidgets.QTextEdit(self.centralwidget)
        self.DNA_Sequnce_input.setObjectName("DNA_Sequnce_input")
        self.gridLayout.addWidget(self.DNA_Sequnce_input, 1, 0, 1, 3)
        self.textEdit_additional_text = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_additional_text.setObjectName("textEdit_additional_text")
        self.gridLayout.addWidget(self.textEdit_additional_text, 4, 0, 1, 1)
        AL_Codon_Counter.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(AL_Codon_Counter)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 879, 21))
        self.menubar.setObjectName("menubar")

        self.menuSelect_File = QtWidgets.QMenu(self.menubar)
        self.menuSelect_File.setStyleSheet("""
                    QMenu {
                        background-color: #3498db; /* Couleur d'arrière-plan */
                        color: white; /* Couleur du texte */
                    }
                    QMenu::item:selected { /* Style des éléments sélectionnés */
                        background-color: #2980b9;
                    }
                """)

        self.menuSave_doc = QtWidgets.QMenu(self.menuSelect_File)
        self.menuSave_doc.setObjectName("menuSave_doc")
        self.menu_INFO = QtWidgets.QMenu(self.menubar)
        self.menu_INFO.setStyleSheet("""
            QMenu {
                background-color: #3498db;  // Background color
                color: white;  // Text color
            }
            QMenu::item:selected {  // Style for selected items
                background-color: #2980b9;
            }
        """)
        self.menu_INFO.setObjectName("?")

        AL_Codon_Counter.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(AL_Codon_Counter)
        self.statusbar.setObjectName("statusbar")
        AL_Codon_Counter.setStatusBar(self.statusbar)
        self.actionsearch_Path = QtWidgets.QAction(AL_Codon_Counter)

        self.actionAbout = QtWidgets.QAction("About", AL_Codon_Counter)
        self.actionAbout.setObjectName("actionAbout")

        self.actionAbout = QtWidgets.QAction(AL_Codon_Counter)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(9)
        self.actionsearch_Path.setFont(font)
        self.actionsearch_Path.setObjectName("actionsearch_Path")
        self.actionrename_file = QtWidgets.QAction(AL_Codon_Counter)
        self.actionrename_file.setObjectName("actionrename_file")
        self.actionSelect_Mlutiple_txt_Fasta_Files = QtWidgets.QAction(AL_Codon_Counter)
        self.actionSelect_Mlutiple_txt_Fasta_Files.setObjectName("actionSelect_Mlutiple_txt_Fasta_Files")
        self.actionAs_txt = QtWidgets.QAction(AL_Codon_Counter)
        self.actionAs_txt.setObjectName("actionAs_txt")
        self.actionAs_docx = QtWidgets.QAction(AL_Codon_Counter)
        self.actionAs_docx.setObjectName("actionAs_docx")
        self.menuSave_doc.addAction(self.actionAs_txt)
        self.menuSave_doc.addAction(self.actionAs_docx)
        self.menuSelect_File.addAction(self.actionsearch_Path)
        self.menuSelect_File.addAction(self.actionSelect_Mlutiple_txt_Fasta_Files)
        self.menuSelect_File.addAction(self.menuSave_doc.menuAction())
        self.menuSelect_File.addAction(self.actionrename_file)

        # Add menus to the menu bar in the correct order
        self.menubar.addAction(self.menuSelect_File.menuAction())

        self.menubar.addAction(self.menu_INFO.menuAction())

        self.menu_INFO.addAction(self.actionAbout)

        self.DNA_Sequnce_output.setStyleSheet("background-color: white;")
        self.textEdit_additional_text.setStyleSheet("background-color: white;")
        self.DNA_Sequnce_input.setStyleSheet("background-color: white;")
        self.textEdit_results.setStyleSheet("background-color: white;")
        self.textEdit_results_TAIL.setStyleSheet("background-color: white;")
        self.list_Sequence_titles_comboBox.setStyleSheet("background-color: white;")
        self.liste_pentamers_comboBox_5.setStyleSheet("background-color: white;")
        self.liste_pentamers_comboBox.setStyleSheet("background-color: white;")
        self.centralwidget.setStyleSheet("background-color: rgb(211, 229, 245);")
        AL_Codon_Counter.setStyleSheet("""
            * {
                font-size: 11pt;
            }
        """)

        self.pushButton_Next_If_Doc_or_multi_Files.setStyleSheet("QPushButton#pushButton_Next_If_Doc_or_multi_Files{\n"
                                                                 "background-color: rgb(43, 17, 242);\n"
                                                                 "color: rgb(255, 255, 255);\n"
                                                                 "border-radius:5px;\n"
                                                                 "text-align: center;\n"
                                                                 "padding-right: 30px;\n"
                                                                 "}\n"
                                                                 "QPushButton#pushButton_Next_If_Doc_or_multi_Files"
                                                                 ":hover{\n "
                                                                 "background-color: qlineargradient(spread:pad, x1:0, "
                                                                 "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(50, 223, "
                                                                 "111, 119), stop:1 rgba(35,81,184,226));\n "
                                                                 "}\n"
                                                                 "QPushButton#pushButton_Next_If_Doc_or_multi_Files"
                                                                 ":pressed{\n "
                                                                 "padding-left:5px;\n"
                                                                 "padding-top:5px;\n"
                                                                 "background-color: qlineargradient(spread:pad, x1:0, "
                                                                 "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(35,81,184,"
                                                                 "226), stop:1  rgba(50, 223, 111, 119));\n "
                                                                 "}")
        iconclip = QtGui.QIcon()
        iconclip.addPixmap(QtGui.QPixmap(resource_path("features/arrow-right-circle.svg")), QtGui.QIcon.Normal,
                           QtGui.QIcon.Off)
        self.pushButton_Next_If_Doc_or_multi_Files.setIcon(iconclip)
        self.pushButton_Next_If_Doc_or_multi_Files.setFixedHeight(30)

        self.pushButton_Create_ALL_in_One.setStyleSheet("QPushButton#pushButton_Create_ALL_in_One{\n"
                                                        "background-color: rgb(43, 17, 242);\n"
                                                        "color: rgb(255, 255, 255);\n"
                                                        "border-radius:5px;\n"
                                                        "text-align: center;\n"
                                                        "padding-right: 30px;\n"
                                                        "}\n"
                                                        "QPushButton#pushButton_Create_ALL_in_One:hover{\n"
                                                        "background-color: qlineargradient(spread:pad, x1:0, "
                                                        "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(50, 223, 111, 119), "
                                                        "stop:1 rgba(35,81,184,226));\n "
                                                        "}\n"
                                                        "QPushButton#pushButton_Create_ALL_in_One:pressed{\n"
                                                        "padding-left:5px;\n"
                                                        "padding-top:5px;\n"
                                                        "background-color: qlineargradient(spread:pad, x1:0, "
                                                        "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(35,81,184,226), "
                                                        "stop:1  rgba(50, 223, 111, 119));\n "

                                                        "}")
        iconclip = QtGui.QIcon()
        iconclip.addPixmap(QtGui.QPixmap(resource_path("features/filter.svg")), QtGui.QIcon.Normal,
                           QtGui.QIcon.Off)
        self.pushButton_Create_ALL_in_One.setIcon(iconclip)

        self.pushButton_Create_ALL_in_One.setFixedHeight(30)

        self.pushButton_renamefile.setStyleSheet("QPushButton#pushButton_renamefile{\n"
                                                 "background-color: rgb(43, 17, 242);\n"
                                                 "color: rgb(255, 255, 255);\n"
                                                 "border-radius:5px;\n"
                                                 "text-align: center;\n"
                                                 "padding-right: 30px;\n"
                                                 "}\n"
                                                 "QPushButton#pushButton_renamefile:hover{\n"
                                                 "background-color: qlineargradient(spread:pad, x1:0, y1:0.505682, "
                                                 "x2:1, y2:0.477, stop:0 rgba(50, 223, 111, 119), stop:1 rgba(35,81,"
                                                 "184,226));\n "
                                                 "}\n"
                                                 "QPushButton#pushButton_renamefile:pressed{\n"
                                                 "padding-left:5px;\n"
                                                 "padding-top:5px;\n"
                                                 "background-color: qlineargradient(spread:pad, x1:0, y1:0.505682, "
                                                 "x2:1, y2:0.477, stop:0 rgba(35,81,184,226), stop:1  rgba(50, 223, "
                                                 "111, 119));\n "

                                                 "}")
        iconclip = QtGui.QIcon()
        iconclip.addPixmap(QtGui.QPixmap(resource_path("features/edit-3.svg")), QtGui.QIcon.Normal,
                           QtGui.QIcon.Off)
        self.pushButton_renamefile.setIcon(iconclip)

        self.pushButton_renamefile.setFixedHeight(30)

        self.pushButton_generate_Docx_seq.setStyleSheet("QPushButton#pushButton_generate_Docx_seq{\n"
                                                        "background-color: rgb(43, 17, 242);\n"
                                                        "color: rgb(255, 255, 255);\n"
                                                        "border-radius:5px;\n"
                                                        "text-align: center;\n"
                                                        "padding-right: 30px;\n"
                                                        "}\n"
                                                        "QPushButton#pushButton_generate_Docx_seq:hover{\n"
                                                        "background-color: qlineargradient(spread:pad, x1:0, "
                                                        "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(50, 223, 111, 119), "
                                                        "stop:1 rgba(35,81,184,226));\n "
                                                        "}\n"
                                                        "QPushButton#pushButton_generate_Docx_seq:pressed{\n"
                                                        "padding-left:5px;\n"
                                                        "padding-top:5px;\n"
                                                        "background-color: qlineargradient(spread:pad, x1:0, "
                                                        "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(35,81,184,226), "
                                                        "stop:1  rgba(50, 223, 111, 119));\n "

                                                        "}")
        iconclip = QtGui.QIcon()
        iconclip.addPixmap(QtGui.QPixmap(resource_path("features/file.svg")), QtGui.QIcon.Normal,
                           QtGui.QIcon.Off)
        self.pushButton_generate_Docx_seq.setIcon(iconclip)

        self.pushButton_generate_Docx_seq.setFixedHeight(30)

        self.pushButton_generate_Txt_seq.setStyleSheet("QPushButton#pushButton_generate_Txt_seq{\n"
                                                       "background-color: rgb(43, 17, 242);\n"
                                                       "color: rgb(255, 255, 255);\n"
                                                       "border-radius:5px;\n"
                                                       "text-align: center;\n"
                                                       "padding-right: 30px;\n"
                                                       "}\n"
                                                       "QPushButton#pushButton_generate_Txt_seq:hover{\n"
                                                       "background-color: qlineargradient(spread:pad, x1:0, "
                                                       "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(50, 223, 111, 119), "
                                                       "stop:1 rgba(35,81,184,226));\n "
                                                       "}\n"
                                                       "QPushButton#pushButton_generate_Txt_seq:pressed{\n"
                                                       "padding-left:5px;\n"
                                                       "padding-top:5px;\n"
                                                       "background-color: qlineargradient(spread:pad, x1:0, "
                                                       "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(35,81,184,226), "
                                                       "stop:1  rgba(50, 223, 111, 119));\n "

                                                       "}")
        iconclip = QtGui.QIcon()
        iconclip.addPixmap(QtGui.QPixmap(resource_path("features/file-text.svg")), QtGui.QIcon.Normal,
                           QtGui.QIcon.Off)
        self.pushButton_generate_Txt_seq.setIcon(iconclip)

        self.pushButton_generate_Txt_seq.setFixedHeight(30)

        self.pushButton_launch_counter.setStyleSheet("QPushButton#pushButton_launch_counter{\n"
                                                     "background-color: rgb(43, 17, 242);\n"
                                                     "color: rgb(255, 255, 255);\n"
                                                     "border-radius:5px;\n"
                                                     "text-align: center;\n"
                                                     "padding-right: 30px;\n"
                                                     "}\n"
                                                     "QPushButton#pushButton_launch_counter:hover{\n"
                                                     "background-color: qlineargradient(spread:pad, x1:0, "
                                                     "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(50, 223, 111, 119), "
                                                     "stop:1 rgba(35,81,184,226));\n "
                                                     "}\n"
                                                     "QPushButton#pushButton_launch_counter:pressed{\n"
                                                     "padding-left:5px;\n"
                                                     "padding-top:5px;\n"
                                                     "background-color: qlineargradient(spread:pad, x1:0, "
                                                     "y1:0.505682, x2:1, y2:0.477, stop:0 rgba(35,81,184,226), "
                                                     "stop:1  rgba(50, 223, 111, 119));\n "

                                                     "}")
        iconclip = QtGui.QIcon()
        iconclip.addPixmap(QtGui.QPixmap(resource_path("features/settings.svg")), QtGui.QIcon.Normal,
                           QtGui.QIcon.Off)
        self.pushButton_launch_counter.setIcon(iconclip)

        self.pushButton_launch_counter.setFixedHeight(30)

        logo_label = QtWidgets.QLabel()
        icon_path_logo = resource_path("JPG_PNG IMAGES/al-codon-high-resolution-logo-transparent.png")
        logo_pixmap = QtGui.QPixmap(icon_path_logo)
        logo_pixmap = logo_pixmap.scaled(100, 100, QtCore.Qt.KeepAspectRatio)
        logo_label.setPixmap(logo_pixmap)
        container_widget = QtWidgets.QWidget(self.centralwidget)
        container_layout = QtWidgets.QVBoxLayout(container_widget)
        container_layout.addWidget(logo_label, alignment=QtCore.Qt.AlignCenter)
        self.gridLayout.addWidget(container_widget, 4, 1, 1, 2)

        # les listes déroulantes
        # Style for QComboBox
        combobox_style_template = """
                QComboBox#{} {{
                    border: 1px solid gray;
                    border-radius: 3px;
                    padding: 1px 18px 1px 3px;
                    min-width: 6em;
                    font-size: 11pt;
                    background-color: white;
                }}

                QComboBox#{}::drop-down {{
                    subcontrol-origin: padding;
                    subcontrol-position: top right;
                    width: 30px;  /* Increased width to make the drop-down area larger */
                    border-left-width: 1px;
                    border-left-color: darkgray;
                    border-left-style: solid;
                    border-top-right-radius: 3px;
                    border-bottom-right-radius: 3px;
                }}
        
                QComboBox#{}::down-arrow {{
                    image: url(features/down-chevron-blue.svg);
                    width: 12px;  /* Adjust width of the arrow */
                    height: 12px;  /* Adjust height of the arrow */
                }}

                QComboBox#{} QAbstractItemView {{
                    border: 1px solid gray;
                    selection-background-color: lightgray;
                }}
                """

        # Apply style to each QComboBox
        self.liste_pentamers_comboBox.setStyleSheet(combobox_style_template.format(
            "liste_pentamers_comboBox", "liste_pentamers_comboBox", "liste_pentamers_comboBox",
            "liste_pentamers_comboBox"))

        self.liste_pentamers_comboBox_5.setStyleSheet(combobox_style_template.format(
            "liste_pentamers_comboBox_5", "liste_pentamers_comboBox_5", "liste_pentamers_comboBox_5",
            "liste_pentamers_comboBox_5"))

        self.list_Sequence_titles_comboBox.setStyleSheet(combobox_style_template.format(
            "list_Sequence_titles_comboBox", "list_Sequence_titles_comboBox", "list_Sequence_titles_comboBox",
            "list_Sequence_titles_comboBox"))

        self.retranslateUi(AL_Codon_Counter)
        QtCore.QMetaObject.connectSlotsByName(AL_Codon_Counter)

        pass

    def retranslateUi(self, AL_Codon_Counter):
        _translate = QtCore.QCoreApplication.translate
        AL_Codon_Counter.setWindowTitle(_translate("AL_Codon_Counter", "AL_Codon_Counter"))
        self.label.setText(_translate("AL_Codon_Counter", "Enter the Sequence:"))
        self.label_5.setText(_translate("AL_Codon_Counter",
                                        "<html><head/><body><p><span style=\" font-weight:600; color:#0000ff;\">If "
                                        ".Docx or Multiple File chosen</span></p></body></html>"))
        self.pushButton_renamefile.setText(_translate("AL_Codon_Counter", "    Rename File"))
        self.pushButton_launch_counter.setText(_translate("AL_Codon_Counter", "    Launch Counter"))
        self.label_9.setText(_translate("AL_Codon_Counter", "Additional Info"))
        self.label_3.setText(_translate("AL_Codon_Counter", "Colored Occurrences:"))
        self.pushButton_Create_ALL_in_One.setText(_translate("AL_Codon_Counter", "    Create ALL in One"))
        self.pushButton_generate_Docx_seq.setText(_translate("AL_Codon_Counter", "    Generate .Docx file"))
        self.label_8.setText(_translate("AL_Codon_Counter",
                                        "<html><head/><body><p><span style=\" font-weight:600; "
                                        "color:#0000ff;\">Generate All Chosen Files in One "
                                        "File</span></p></body></html>"))
        self.pushButton_Next_If_Doc_or_multi_Files.setText(_translate("AL_Codon_Counter", "    Next"))
        self.pushButton_generate_Txt_seq.setText(_translate("AL_Codon_Counter", "    Generate .Txt file"))
        self.radio_red_Button.setText(_translate("AL_Codon_Counter", "Red Pentamers"))
        self.colored_radio_Button.setText(_translate("AL_Codon_Counter", "Colored Pentamers"))
        self.colored_radio_Button_tail.setText(_translate("AL_Codon_Counter", "Colored Pentamers Tail"))
        self.red_radio_Button_tail.setText(_translate("AL_Codon_Counter", "Red Pentamers Tail"))
        self.label_7.setText(_translate("AL_Codon_Counter",
                                        "<html><head/><body><p><span style=\" font-weight:600; "
                                        "color:#0000ff;\">Select To view File </span></p><p><span style=\" "
                                        "font-weight:600; color:#0000ff;\">from the previously chosen "
                                        "files</span></p></body></html>"))
        self.label_4.setText(_translate("AL_Codon_Counter", "Selected Document:"))
        self.label_10.setText(_translate("AL_Codon_Counter", "Occurrences/Results Tail-- 5 Pentamers :"))
        self.label_2.setText(_translate("AL_Codon_Counter", "Occurrences/Results -- 9 Pentamers:"))
        self.label_6.setText(_translate("AL_Codon_Counter", "-> Select Pentamer from 9 Pentamers"))
        self.label_5_Pentamers.setText(_translate("AL_Codon_Counter", "  -> Select Pentamer from the Tail"))

        self.menuSelect_File.setTitle(_translate("AL_Codon_Counter", "Select File "))
        self.menuSave_doc.setTitle(_translate("AL_Codon_Counter", "Save doc"))

        self.menu_INFO.setTitle(_translate("AL_Codon_Counter", "?"))

        self.actionAbout.setText(_translate("AL_Codon_Counter", "About"))

        self.actionsearch_Path.setText(_translate("AL_Codon_Counter", "Search Path"))
        self.actionrename_file.setText(_translate("AL_Codon_Counter", "    Rename file"))
        self.actionSelect_Mlutiple_txt_Fasta_Files.setText(
            _translate("AL_Codon_Counter", "Select Multiple (.txt/.Fasta) Files "))
        self.actionAs_txt.setText(_translate("AL_Codon_Counter", "As .txt"))
        self.actionAs_docx.setText(_translate("AL_Codon_Counter", "As .docx"))
        pass


class ALCodonCounterMainWindow(QMainWindow):
    def __init__(self, parent=None):
        super(ALCodonCounterMainWindow, self).__init__(parent)
        self.results = None
        self.results_TAIL = None
        self.result_file = None
        self.ui = Ui_AL_Codon_Counter()
        self.ui.setupUi(self)
        self.directory_changed = 0
        # -----------------------------------

        # -----------------------------------
        # Setup View_Red_or_Colored_Selected_Sequence window Ui_MainWindow_3()
        # Manager for shared resources
        self.manager = Manager()
        self.results_queue = self.manager.Queue()
        # Start the result listener thread
        self.result_listener = ResultListenerThread(self.results_queue)
        self.result_listener.data_processed.connect(self.update_gui)
        self.result_listener.start()

        # Check for results periodically
        self.timer = QTimer()
        self.timer.timeout.connect(self.check_results)

        self.counter = AL_CODON_Counter()
        self.counter_docx = Read_Doc_and_generate_Txt()
        self.fname = ""
        self.combo_titles = []
        self.sequence_data = {}
        self.current_index = 0
        self.titles = []
        self.cleaned_sequences = []
        self.additional_text = []
        # Define the strings to search for
        self.search_strings = ["ATTCA", "TTCAA", "TCAAG", "CAAGA", "AAGAT", "AGATG", "GATGA", "ATGAA", "TGAAT"]

        # Map each search string to a specific color
        self.color_mapping = {
            "ATTCA": "red",
            "TTCAA": "blue",
            "TCAAG": "green",
            "CAAGA": "purple",
            "AAGAT": "orange",
            "AGATG": "pink",
            "GATGA": "brown",
            "ATGAA": "teal",
            "TGAAT": "gray"
        }

        # ------------------------------------------------------------------------------------
        # In Case There is Nothing Selected YET Handle an ERROR Message
        # Generates a Message when the pushButton_Next_If_Doc_or_multi_Files
        # is clicked and Input sequence is EMPTY
        self.ui.pushButton_Next_If_Doc_or_multi_Files.clicked.connect(self.handle_button_click)
        # Generates A Message when Create_ALL_in_One .docx file Button is clicked and Input sequence is EMPTY
        self.ui.pushButton_Create_ALL_in_One.clicked.connect(self.handle_button_click)
        # -------------------------------------------------------------------------------------

        # Connect the actionsearch_Path action to the clicker method
        self.ui.actionsearch_Path.triggered.connect(self.clicker)

        # Connect the actionAs_txt and actionAs_docx actions to the generate_Text_file /generate_Doc_file methods
        # respectively
        self.ui.actionAs_txt.triggered.connect(self.generate_Text_file)
        self.ui.actionAs_docx.triggered.connect(self.generate_Doc_file)
        # Connect the actionSelect_Mlutiple_txt_Fasta_Files action
        self.ui.actionSelect_Mlutiple_txt_Fasta_Files.triggered.connect(self.clicker_Multi_Files)
        # Connect  the self.actionAbout  action

        # Connect the triggered signal to the slot
        try:
            self.ui.actionAbout.triggered.connect(self.showAboutDialog)
        except Exception as e:
            QMessageBox.critical(self, "Error showAboutDialog", f'Failed to connect action: {e}')

        # to  Select Multiple File either .Txt/.Fasta or
        # combined Connect the Rename File actions to the rename_file method
        self.ui.actionrename_file.triggered.connect(self.rename_file)

        # ---- Renamme the file
        self.ui.pushButton_renamefile.clicked.connect(self.rename_file)
        # ---- generate .Txt file
        self.ui.pushButton_generate_Txt_seq.clicked.connect(self.generate_Text_file)
        # ---- generate .docx file
        self.ui.pushButton_generate_Docx_seq.clicked.connect(self.generate_Doc_file)
        # ---- Launch find_and_count_with_coloring & Display_results

        self.ui.pushButton_launch_counter.clicked.connect(self.handle_launch)
        # Make ComboBox liste_pentamers Clickable
        self.ui.liste_pentamers_comboBox.activated.connect(self.openWindow)
        self.ui.liste_pentamers_comboBox_5.activated.connect(self.openWindow_TAIL)

        # Add a liste_pentamers
        my_toppings = ["ALL", "ATTCA", "TTCAA", "TCAAG", "CAAGA", "AAGAT", "AGATG", "GATGA", "ATGAA", "TGAAT"]
        self.ui.liste_pentamers_comboBox.addItems(my_toppings)
        my_toppings_TAIL = ["ALL", "AGATG", "GATGA", "GTGGC", "TGGCC", "GGCCT"]
        self.ui.liste_pentamers_comboBox_5.addItems(my_toppings_TAIL)
        # Set Button States
        self.ui.radio_red_Button.toggled.connect(lambda: self.openWindow_red_colored_seq(self.ui.radio_red_Button))
        self.ui.colored_radio_Button.toggled.connect(
            lambda: self.openWindow_red_colored_seq(self.ui.colored_radio_Button))
        self.ui.colored_radio_Button_tail.toggled.connect(
            lambda: self.openWindow_red_colored_seq(self.ui.colored_radio_Button_tail))
        self.ui.red_radio_Button_tail.toggled.connect(
            lambda: self.openWindow_red_colored_seq(self.ui.red_radio_Button_tail))

    # Show the ABOUT the AL CODON Information window
    def showAboutDialog(self):

        dialog = QtWidgets.QMessageBox(self)
        dialog.setWindowTitle("About AL Codon Counter")
        # Allows the use of HTML
        dialog.setTextFormat(QtCore.Qt.RichText)

        # Configuring dialog content including logo
        about_text = (
            "<h3>AL Codon Counter</h3>"
            "<p><b>Version:</b> 1.0</p>"
            "<p><b>Year:</b> 2024</p>"
            "<p><b>Developer:</b> Houssem Ben Khalfallah</p>"
            "<p>PhD, Bureau AGEIS : 332 (3ème étage) Faculté de médecine - Bâtiment Jean Roget 38706 La Tronche, "
            "ED-ISCE, Université Grenoble Alpes</p>"
            "<p><b>Email:</b> <a href='mailto:houssembenkhalfallah@gmail.com'>houssembenkhalfallah@gmail.com</a></p>"
            "<p><b>Alternate Email:</b> <a href='mailto:Houssem.Ben-Khalfallah@univ-grenoble-alpes.fr'>Houssem.Ben"
            "-Khalfallah@univ-grenoble-alpes.fr</a></p> "
            "<p><b>Research Supervisor:</b> Jacques Demongeot, Professor Emeritus at Université Grenoble Alpes</p>"
            "<p>This software is a research tool developed for data mining and analysis in the field of genetics. It "
            "provides efficient and accurate data handling capabilities to enhance research outcomes.</p>"
            "<p><b>© 2024 Houssem Ben Khalfallah. All rights reserved.</b></p>"
        )

        # Construct the path to the image
        image_path = resource_path("JPG_PNG IMAGES/logo.png")

        # Load the image
        logo = QtGui.QPixmap(image_path)

        dialog.setIconPixmap(logo.scaled(64, 64, QtCore.Qt.KeepAspectRatio))

        dialog.setText(about_text)
        dialog.exec_()

    # Function to convert RGB tuple to docx RGBColor
    def rgb_tuple_to_rgbcolor(self, rgb_tuple):
        return RGBColor(int(rgb_tuple[0] * 255), int(rgb_tuple[1] * 255), int(rgb_tuple[2] * 255))

    def add_colored_heading(self, document, text, level, rgb_color):
        """ Adds a heading with specified color to a Word document. """
        heading = document.add_heading(level=level)
        run = heading.add_run(text)
        run.font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])
        # Optionally set the font size (Houssem decided THAT)
        run.font.size = Pt(16)

    def find_and_print_remainder(self, sequence, target_pattern):
        """
        Searches for a 10-character pattern from the end of the sequence, moving backwards.
        Once the pattern is found, prints the preceding part of the sequence and returns it.

        Args:
        sequence (str): The DNA sequence to search within.
        target_pattern (str): The 10-character string to find in the sequence.

        Returns:
        str: The remaining text before the found pattern or an empty string if no match is found.
        """
        # Ensure the target pattern is exactly 10 characters
        if len(target_pattern) != 10:
            QMessageBox.information(self, " Ensure the target pattern is exactly 10 characters",
                                    "The target pattern must be exactly 10 characters long.")

            return ""

        # Reverse the sequence to simplify searching from the end
        reversed_sequence = sequence[::-1]

        # Iterate over the reversed sequence
        for i in range(len(reversed_sequence) - 9):
            # Extract a substring of 10 characters starting from index i
            substring = reversed_sequence[i:i + 10]

            # Check if the extracted substring matches the reversed target pattern
            # Compare against reversed pattern
            if substring == target_pattern[::-1]:
                # Calculate the starting point of the remaining sequence
                start_index = len(sequence) - i - 10
                # Print the remaining sequence up to the match
                remaining_text = sequence[:start_index]

                return remaining_text
        # If No match Found --> RETURN the input Sequence
        QMessageBox.warning(self, "Error", "No match found.")
        return sequence

    def generate_Doc_Create_ALL_in_One(self):
        try:
            # Check if the DNA_Sequnce_intput is empty
            if not self.ui.DNA_Sequnce_input.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence input is empty.\n"
                                                         "It seems That there is No File selected.")
                return

            # Check if the DNA_Sequnce_output is empty
            if not self.ui.DNA_Sequnce_output.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence output 'Colored Occurrences' is empty.\n"
                                                         "It seems that You did Not Click On the 'Launch Counter' "
                                                         "Button "
                                                         "\n "
                                                         "Before Generating a .Docx ALL Sequences in ONE File.")
                return

            global sequence_text_cutted
            color_mapping = {
                "ATTCA": (255, 0, 0),
                "TTCAA": (0, 0, 255),
                "TCAAG": (0, 128, 0),
                "CAAGA": (128, 0, 128),
                "AAGAT": (255, 165, 0),
                "AGATG": (255, 192, 203),
                "GATGA": (165, 42, 42),
                "ATGAA": (0, 128, 128),
                "TGAAT": (128, 128, 128)
            }
            # Initial path before showing the dialog
            # Open a file dialog to select the location and name of the output file
            # Create a QFileDialog instance
            dialog = QFileDialog(self, "Save Sequence", "")
            dialog.setAcceptMode(QFileDialog.AcceptSave)
            dialog.setNameFilters(["Word Files (*.docx)", "All Files (*)"])
            # Connect the signal to monitor directory changes
            self.directory_changed = 0
            dialog.directoryEntered.connect(self.on_directory_changed)
            # Execute the dialog and get the selected filename
            if dialog.exec_() == QFileDialog.Accepted:
                filename = dialog.selectedFiles()[0]
                if not filename.strip():
                    QMessageBox.warning(self, "No Selection", "No file selected.")
                    return
            else:
                QMessageBox.warning(self, "Cancelled", "The ALL in ONE Doc File generating cancelled.")
                return

            # ------------------------------------------------------------------------------------------------

            if filename:
                try:
                    # If the user selected a file, write the sequence text to it
                    # Create a new Word document
                    doc = Document()
                    # Loop in self.sequence_data to Get sequences which are the input cleaned_sequences
                    for title, sequence in self.sequence_data.items():

                        # Get sequence Colored <--results_sequence[3]
                        results_sequence = self.counter.find_and_count_with_coloring(sequence, self.search_strings,
                                                                                     color_mapping)
                        results_sequence_TAIL = self.counter.find_and_count_with_coloring_TAIL(sequence)
                        total_occurrences = results_sequence[1]
                        Colored_results_sequence = results_sequence[3]

                        # Get Occurrences/Results for sequence
                        results_summary = self.counter.Display_results(results_sequence)
                        results_summary_TAIL = self.counter.Display_results_TAIL(results_sequence_TAIL)
                        # Get corresponding Title
                        results_Title = title

                        # The main idea here is to get the NUMBER of ONLY COLORED pentamers to use it
                        # later to break the loop of for match in regex.finditer(sequence_text) to get the right
                        # sequence without ADDING Suppose 'Colored_results_sequence' is your HTML containing
                        # colored pentamers
                        soup_red = BeautifulSoup(Colored_results_sequence, "html.parser")
                        # Revert All the Colored Pentamers to RED
                        # Assuming color is applied directly in style, adjust to match actual output
                        colored_text_elements = soup_red.find_all(style=lambda value: value and 'color' in value)

                        # Define pentamers and initialize count dictionary
                        pentamers = ["ATTCA", "TTCAA", "TCAAG", "CAAGA", "AAGAT", "AGATG", "GATGA", "ATGAA", "TGAAT"]
                        pentamer_counts = {pentamer: 0 for pentamer in pentamers}

                        # Count occurrences of each pentamer in elements with the specified color
                        for element in colored_text_elements:
                            text = element.get_text()
                            for pentamer in pentamers:
                                pentamer_counts[pentamer] += text.count(pentamer)

                        # Add Heading in the Doc
                        self.add_colored_heading(doc, results_Title, level=1, rgb_color=(255, 0, 0))
                        doc.add_heading('DNA Sequence', level=1)
                        # ---------------------------------------------------------------------------------------------------------

                        # Parse elements in the list Colored_results_sequence
                        soup = BeautifulSoup(Colored_results_sequence, 'html.parser')
                        paragraph = doc.add_paragraph()
                        previous_text_ended_with_newline = False  # Indicates if the previous text ends with a newline
                        for element in soup:
                            if element.name == 'span':
                                style = element.get('style')
                                # Use a regular expression to extract RGB values
                                color_match = re.search(r'\((\d+), (\d+), (\d+)\)', style)
                                if color_match:
                                    r, g, b = map(int, color_match.groups())
                                    # Add colored text to the paragraph with the appropriate color
                                    run = paragraph.add_run(element.text)
                                    run.font.color.rgb = RGBColor(r, g, b)
                                    if previous_text_ended_with_newline:
                                        # If the previous text ended with a newline, remove it
                                        paragraph.runs[-2].text = paragraph.runs[-2].text.rstrip('\n')
                                    previous_text_ended_with_newline = False
                            else:
                                # Add black text to the paragraph
                                paragraph.add_run(element)
                                if element.endswith('\n'):
                                    previous_text_ended_with_newline = True
                                else:
                                    previous_text_ended_with_newline = False

                        # ----------------------------------------------------------------------------------------
                        doc.add_heading('Occurrences/Results 9 Pentamers', level=1)
                        doc.add_paragraph(results_summary)
                        doc.add_heading('Occurrences/Results Tail 5 Pentamers', level=1)
                        doc.add_paragraph(results_summary_TAIL)

                    # Get the directory path If the user didn't choose a directory, use the "Documents" directory
                    # FOR Windows under C:\Users\<username>\Documents If the user didn't choose a directory
                    # FOR Mac under \Users\<username>\Documents If the user didn't choose a directory
                    # FOR Linux under \Users\<username>\Documents If the user didn't choose a directory
                    if self.directory_changed > 1:
                        # If the user chose a directory
                        full_path = filename
                    else:
                        # If the user did not choose a directory
                        folder_path = os.path.join(os.path.expanduser("~"), "Documents", "Generated ALL in One DOCX")
                        # Ensure the directory exists
                        os.makedirs(folder_path, exist_ok=True)
                        full_path = os.path.join(folder_path, os.path.basename(filename))
                    # Get the file name
                    full_path = self.save_document_in_folder(doc, full_path)

                    QMessageBox.information(self, "Success", f"Successfully saved the sequence to {full_path}")
                except IOError as e:
                    QMessageBox.critical(self, "Error", f"An error occurred while saving the file: {e}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An unexpected error occurred: {e}")

    # View the Change of directory PAth
    def on_directory_changed(self, directory):
        print('Directory changed:', directory)
        self.directory_changed += 1
        print('self.directory_changed = ', self.directory_changed)

    # The save_document_in_folder function checks to ensure that the folder exists and creates it if it doesn't.
    def save_document_in_folder(self, doc, full_path):
        folder_path = os.path.dirname(full_path)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        doc.save(full_path)
        return full_path

    def generate_Text_file(self):
        try:
            # Check if the DNA_Sequnce_intput is empty
            if not self.ui.DNA_Sequnce_input.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence input is empty.\n"
                                                         "It seems That there is No File selected. ")
            else:
                # Check if the DNA_Sequnce_outtput is empty
                if not self.ui.DNA_Sequnce_output.toPlainText().strip():
                    QMessageBox.warning(self, "Input Check", "The Sequence output 'Colored Occurrences' is empty.\n"
                                                             "It seems that You did Not Click On the 'Launch Counter' "
                                                             "Button "
                                                             "\n "
                                                             "Before Generating a .txt File.")
                else:

                    # Retrieve the texts from DNA_Sequnce_output and Occurrences/Results
                    sequence_text = self.ui.DNA_Sequnce_output.toPlainText()
                    Occurrences_Results = self.ui.textEdit_results.toPlainText()
                    Occurrences_Results_TAIL = self.ui.textEdit_results_TAIL.toPlainText()
                    # Retrieve the initial filename or path from the UI widget
                    fname = self.ui.selected_fileshow.text()
                    if fname:
                        # Replace any non-alphanumeric characters with underscores
                        fname = re.sub(r'[^a-zA-Z0-9]', '_', fname)
                    else:
                        fname = ""

                    # Open a file dialog to select the location and name of the output file
                    # Create a QFileDialog instance
                    dialog = QFileDialog(self, "Save Sequence", fname)
                    dialog.setAcceptMode(QFileDialog.AcceptSave)
                    dialog.setNameFilters(["Text Files (*.txt)", "All Files (*)"])
                    # Connect the signal to monitor directory changes
                    self.directory_changed = 0
                    dialog.directoryEntered.connect(self.on_directory_changed)
                    # Execute the dialog and get the selected filename
                    if dialog.exec_() == QFileDialog.Accepted:
                        filename = dialog.selectedFiles()[0]
                        if not filename.strip():
                            QMessageBox.warning(self, "No Selection", "No file selected.")
                            return
                    else:
                        QMessageBox.warning(self, "Cancelled", "Text File generating cancelled.")
                        return

                    # If the user selected a file, write the sequence text to it
                    if filename:
                        # Path of the current file or directory I'm working with
                        # Get the directory path If the user didn't choose a directory, use the "Documents" directory
                        # FOR Windows under C:\Users\<username>\Documents If the user didn't choose a directory
                        # FOR Mac  under \Users\<username>\Documents If the user didn't choose a directory
                        # FOR Linux under \Users\<username>\Documents If the user didn't choose a directory
                        if self.directory_changed > 1:
                            # If the user chose a directory
                            full_path = filename
                        else:
                            # If the user did not choose a directory
                            folder_path = os.path.join(os.path.expanduser("~"), "Documents", "Genereted TXT Files")
                            # Ensure the directory exists
                            os.makedirs(folder_path, exist_ok=True)
                            full_path = os.path.join(folder_path, os.path.basename(filename))

                        # Write the sequence text to the file
                        try:
                            with open(full_path, 'w') as file:
                                file.write(sequence_text)
                                # Add a newline character to move to the next line
                                file.write("\n")
                                file.write("\n")
                                file.write(Occurrences_Results)
                                # Add a newline character to move to the next line
                                file.write("\n")
                                file.write("\n")
                                file.write(Occurrences_Results_TAIL)

                            QMessageBox.information(self, "Success", f"Successfully saved the sequence to {full_path}")
                        except IOError as e:
                            QMessageBox.critical(self, "Error", f"An error occurred while saving the file: {e}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An unexpected error occurred: {e}")

    def generate_Doc_file(self):
        try:
            # Check if the DNA_Sequnce_intput is empty
            if not self.ui.DNA_Sequnce_input.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence input is empty.\n"
                                                         "It seems That there is No File selected.")
                return

            # Check if the DNA_Sequnce_output is empty
            if not self.ui.DNA_Sequnce_output.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence output 'Colored Occurrences' is empty.\n"
                                                         "It seems that You did Not Click On the 'Launch Counter' "
                                                         "Button "
                                                         "\n "
                                                         "Before Generating a .Docx File.")
                return

            # Define DNA sequence to color mapping using matplotlib colors
            color_mapping = {
                "ATTCA": mcolors.to_rgb(mcolors.CSS4_COLORS['red']),
                "TTCAA": mcolors.to_rgb(mcolors.CSS4_COLORS['blue']),
                "TCAAG": mcolors.to_rgb(mcolors.CSS4_COLORS['green']),
                "CAAGA": mcolors.to_rgb(mcolors.CSS4_COLORS['purple']),
                "AAGAT": mcolors.to_rgb(mcolors.CSS4_COLORS['orange']),
                "AGATG": mcolors.to_rgb(mcolors.CSS4_COLORS['magenta']),
                "GATGA": mcolors.to_rgb(mcolors.CSS4_COLORS['brown']),
                "ATGAA": mcolors.to_rgb(mcolors.CSS4_COLORS['teal']),
                "TGAAT": mcolors.to_rgb(mcolors.CSS4_COLORS['darkviolet'])
            }
            # Retrieve the HTML content and results
            html_content = self.ui.DNA_Sequnce_output.toPlainText()
            occurrences_results = self.ui.textEdit_results.toPlainText()
            Occurrences_Results_TAIL = self.ui.textEdit_results_TAIL.toPlainText()
            # Retrieve the initial filename or path from the UI widget
            fname = self.ui.selected_fileshow.text()
            if fname:
                # Replace any non-alphanumeric characters with underscores
                fname = re.sub(r'[^a-zA-Z0-9]', '_', fname)
            else:
                fname = ""

            # Open a file dialog to select the location and name of the output file
            # Create a QFileDialog instance
            dialog = QFileDialog(self, "Save Sequence", fname)
            dialog.setAcceptMode(QFileDialog.AcceptSave)
            dialog.setNameFilters(["Word Files (*.docx)", "All Files (*)"])

            # Connect the signal to monitor directory changes
            self.directory_changed = 0

            dialog.directoryEntered.connect(self.on_directory_changed)

            # Execute the dialog and get the selected filename
            if dialog.exec_() == QFileDialog.Accepted:
                filename = dialog.selectedFiles()[0]
                if not filename.strip():
                    QMessageBox.warning(self, "No Selection", "No file selected.")
                    return
            else:
                QMessageBox.warning(self, "Cancelled", "Doc File generating cancelled.")
                return

            # If the user selected a file, write the sequence text to it
            if filename:
                try:
                    # Parse the HTML content
                    soup = BeautifulSoup(html_content, 'lxml')
                    sequence_text = ''.join([element.get_text() for element in soup.find_all(['span', 'p', 'div'])])

                    # Create a new Word document
                    doc = Document()
                    doc.add_heading('DNA Sequence', level=1)

                    # Start with a blank paragraph
                    paragraph = doc.add_paragraph()

                    # Sort color mapping keys by length in descending order to handle overlapping patterns
                    sorted_keys = sorted(color_mapping.keys(), key=len, reverse=True)

                    # Create a regex pattern to capture all sequences in color mapping
                    pattern = '|'.join(map(re.escape, sorted_keys))
                    regex = re.compile(pattern)

                    # Index to keep track of where I parsed so far
                    last_index = 0

                    # Find all matching sequences
                    for match in regex.finditer(sequence_text):
                        # Add the non-colored portion before the current match
                        if last_index < match.start():
                            paragraph.add_run(sequence_text[last_index:match.start()]).font.color.rgb = RGBColor(0, 0,
                                                                                                                 0)

                        # Add the colored portion (matched sequence)
                        sequence = match.group(0)
                        rgb = color_mapping[sequence]
                        run = paragraph.add_run(sequence)
                        run.font.color.rgb = self.rgb_tuple_to_rgbcolor(rgb)

                        # Update the last index
                        last_index = match.end()

                    # Calculate the trimmed length
                    max_length = len(html_content) + len('DNA Sequence')

                    sequence_text = sequence_text[:max_length]

                    # Add remaining non-colored text
                    if last_index < len(sequence_text):
                        paragraph.add_run(sequence_text[last_index:]).font.color.rgb = RGBColor(0, 0, 0)
                    # Add the occurrences/results
                    doc.add_heading('Occurrences/Results 9 Pentamers', level=1)
                    doc.add_paragraph(occurrences_results)
                    # Add the occurrences/results TAIL
                    doc.add_heading('Occurrences/Results Tail 5 Pentamers', level=1)
                    doc.add_paragraph(Occurrences_Results_TAIL)

                    # Save the document to the selected file
                    # --------------------------------------

                    # Get the directory path If the user didn't choose a directory, use the "Documents" directory
                    # FOR Windows under C:\Users\<username>\Documents If the user didn't choose a directory
                    # FOR Mac  under \Users\<username>\Documents If the user didn't choose a directory
                    # FOR Linux under \Users\<username>\Documents If the user didn't choose a directory
                    if self.directory_changed > 1:
                        # If the user chose a directory
                        full_path = filename
                    else:
                        # If the user did not choose a directory
                        folder_path = os.path.join(os.path.expanduser("~"), "Documents", "Generated DOCX Files")
                        # Ensure the directory exists
                        os.makedirs(folder_path, exist_ok=True)
                        full_path = os.path.join(folder_path, os.path.basename(filename))
                    # Get the file name
                    full_path = self.save_document_in_folder(doc, full_path)

                    QMessageBox.information(self, "Success", f"Successfully saved the sequence to {full_path}")

                except IOError as e:
                    QMessageBox.critical(self, "Error", f"An error occurred while saving the file: {e}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An unexpected error occurred: {e}")

    # Open window to show the Selected Pentamer from the 9 Pentamers
    def openWindow(self):
        # Check if the DNA_Sequnce_input is empty
        if not self.ui.DNA_Sequnce_input.toPlainText().strip():
            QMessageBox.warning(self, "Input Check", "The Sequence Input is empty.\n"
                                                     "It seems that there is No File selected.")
        else:
            # Check if the DNA_Sequnce_outtput is empty
            if not self.ui.DNA_Sequnce_output.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence output 'Colored Occurrences' is empty.\n"
                                                         "It seems that You did Not Click On the 'Launch Counter' "
                                                         "Button "
                                                         "\n "
                                                         "Before Viewing a Selected Pentamer .")
            else:
                # QMessageBox.information(self, "Input Check", "The DNA Sequence Input is not empty.")
                self.window = QtWidgets.QMainWindow()
                self.ui_2 = Ui_MainWindow_2()
                self.ui_2.setupUi(self.window)
                self.window.show()
                # Print the selected pentamer in a highlighted way "red"
                self.clicker_list_pentamers()

    # Open window to show the Selected Pentamer of the 5 Pentamers from te Tail
    def openWindow_TAIL(self):
        # Check if the DNA_Sequnce_input is empty
        if not self.ui.DNA_Sequnce_input.toPlainText().strip():
            QMessageBox.warning(self, "Input Check", "The Sequence Input is empty.\n"
                                                     "It seems that there is No File selected.")
        else:
            # Check if the DNA_Sequnce_outtput is empty
            if not self.ui.DNA_Sequnce_output.toPlainText().strip():
                QMessageBox.warning(self, "Input Check", "The Sequence output 'Colored Occurrences' is empty.\n"
                                                         "It seems that You did Not Click On the 'Launch Counter' "
                                                         "Button "
                                                         "\n "
                                                         "Before Viewing a TAIL Selected Pentamer .")
            else:
                # QMessageBox.information(self, "Input Check", "The DNA Sequence Input is not empty.")
                self.window = QtWidgets.QMainWindow()
                self.ui_2 = Ui_MainWindow_2()
                self.ui_2.setupUi(self.window)
                self.window.show()
                # Print the selected pentamer in a highlighted way "red"
                self.clicker_list_pentamers_TAIL()

    def on_subwindow_close(self):
        """ Custom close event handler of window_red_colored_seq. """
        # Uncheck the radio button
        self.ui.radio_red_Button.setChecked(False)
        self.ui.colored_radio_Button.setChecked(False)

    def openWindow_red_colored_seq(self, radio_Button):
        # Check if the DNA_Sequnce_input is empty
        if not self.ui.DNA_Sequnce_input.toPlainText().strip():
            QMessageBox.warning(self, "Input Check", "The Sequence Input is empty.\n"
                                                     "It seems that there is No File selected.")
        else:
            # Ensure UI and window are initialized only once and reused
            if not hasattr(self, 'ui_3'):
                self.ui_3 = Ui_MainWindow_3()
                self.window_red_colored_seq = QtWidgets.QMainWindow()
                self.ui_3.setupUi(self.window_red_colored_seq)

            # Display the window regardless of which button is selected
            self.window_red_colored_seq.show()

            # Conditionally set the HTML content based on the radio button selection
            if radio_Button == self.ui.colored_radio_Button:
                self.ui_3.textEdit_Selected_Sequence.clear()
                # Set HTML content only when the colored radio button is selected
                text_input = self.ui.DNA_Sequnce_input.toPlainText()
                colored_text = self.counter.find_and_count_to_Colored(text_input)
                self.ui_3.textEdit_Selected_Sequence.setHtml(colored_text)
            if radio_Button == self.ui.radio_red_Button:
                # Optionally handle the red button selection,clearing or setting different content
                # Clear previous content or handle differently
                self.ui_3.textEdit_Selected_Sequence.clear()
                text_input_red = self.ui.DNA_Sequnce_input.toPlainText()
                red_text = self.counter.find_and_count_to_RED(text_input_red)
                self.ui_3.textEdit_Selected_Sequence.setHtml(red_text)
            if radio_Button == self.ui.red_radio_Button_tail:
                # Optionally handle the red button selection,clearing or setting different content
                # Clear previous content or handle differently
                self.ui_3.textEdit_Selected_Sequence.clear()
                text_input_red = self.ui.DNA_Sequnce_input.toPlainText()
                red_text = self.counter.find_and_count_to_RED_TAIL(text_input_red)
                self.ui_3.textEdit_Selected_Sequence.setHtml(red_text)
            elif radio_Button == self.ui.colored_radio_Button_tail:
                self.ui_3.textEdit_Selected_Sequence.clear()
                # Set HTML content only when the colored radio button is selected
                text_input = self.ui.DNA_Sequnce_input.toPlainText()
                colored_text = self.counter.find_and_count_to_Colored_TAIL(text_input)
                self.ui_3.textEdit_Selected_Sequence.setHtml(colored_text)

    # clicker_list_pentamers  9 Pentamers
    def clicker_list_pentamers(self):
        selected_pentamer = self.ui.liste_pentamers_comboBox.currentText()
        data_frames = {
            "ATTCA": self.results[8],
            "TTCAA": self.results[9],
            "TCAAG": self.results[10],
            "CAAGA": self.results[11],
            "AAGAT": self.results[12],
            "AGATG": self.results[13],
            "GATGA": self.results[14],
            "ATGAA": self.results[15],
            "TGAAT": self.results[16]
        }

        # Display the result corresponding to the selected pentamer
        if selected_pentamer in data_frames:
            self.ui_2.textEdit_Selected_Pentamer_Occurence.setText(data_frames[selected_pentamer])

    # clicker_list_pentamers_TAIL 5 Pentamers
    def clicker_list_pentamers_TAIL(self):
        selected_pentamer = self.ui.liste_pentamers_comboBox_5.currentText()
        data_frames = {
            "AGATG": self.results_TAIL[8],
            "GATGA": self.results_TAIL[9],
            "GTGGC": self.results_TAIL[10],
            "TGGCC": self.results_TAIL[11],
            "GGCCT": self.results_TAIL[12],
        }

        # Display the result corresponding to the selected pentamer
        if selected_pentamer in data_frames:
            self.ui_2.textEdit_Selected_Pentamer_Occurence.setText(data_frames[selected_pentamer])

    # Select One File .txt/ .Fasta / or .docx
    def clicker(self):

        # delete text in content
        self.ui.DNA_Sequnce_output.clear()
        self.ui.textEdit_results.clear()
        self.ui.textEdit_results_TAIL.clear()
        self.ui.list_Sequence_titles_comboBox.clear()
        # Get the selected file name (fname)
        fname, _ = QFileDialog.getOpenFileName(self, "Open Sequence File For AL-Codon-Counter", "",
                                               "All Files (*);;Text Files (*.txt);;FASTA Files (*.fasta *.fas *.fa "
                                               "*.fna)")

        if fname:
            self.ui.selected_fileshow.setText(fname)
            self.show_ANY_file(fname)

    # Get the Filename without extension
    def get_filename_without_extension(self, file_path):
        # Extract the base name, with extension
        base_name = os.path.basename(file_path)

        # Split the base name by the dot and discard the extension
        file_name_without_extension = os.path.splitext(base_name)[0]

        return file_name_without_extension

    # Select Multiple files .txt / .Fata or combined
    def clicker_Multi_Files(self):
        self.combo_titles = []
        self.sequence_data = {}
        self.titles = []
        self.cleaned_sequences = []
        self.additional_text = []

        # delete text in content
        self.ui.DNA_Sequnce_output.clear()
        self.ui.textEdit_results.clear()
        self.ui.textEdit_results_TAIL.clear()
        self.ui.list_Sequence_titles_comboBox.clear()
        # Get the Multiple selected files names (fname)
        files, _ = QFileDialog.getOpenFileNames(self, "Select Multiple Files either .txt/.Fasta or Combined For "
                                                      "AL-Codon-Counter", "", "Text and FASTA Files (*.txt *.fasta "
                                                                              "*.fas *.fa *.fna)")

        # Check if the user has selected any files
        if files:

            # List comprehension to filter .txt files and get ficount_nucleotide_sequenceslenames without extension

            combined_df = self.counter_docx.process_files_to_dataframe(files)

            # Joining all the text in the 'Content' column into one single string
            combined_content = ''.join(combined_df['Content'])

            self.result_file = self.counter_docx.extract_titles_and_sequences_Multi_files(combined_content)
            # Extract mandatory values first
            text_File = self.result_file[0]
            file_new_title = self.result_file[1]

            self.additional_text = self.result_file[2]

            self.titles = self.result_file[3]
            self.cleaned_sequences = self.result_file[4]

            # Populate combo box and store sequence data
            self.combo_titles = self.titles

            self.sequence_data = dict(zip(self.titles,
                                          self.cleaned_sequences))

            self.ui.DNA_Sequnce_input.setPlainText(self.sequence_data[self.titles[0]])

            self.ui.selected_fileshow.setText(file_new_title)

            self.ui.textEdit_additional_text.setText(self.additional_text[0])

            self.ui.list_Sequence_titles_comboBox.clear()

            # Add a list_Sequence_titles
            self.ui.list_Sequence_titles_comboBox.addItems(self.titles)

            # Make list_Sequence_titles_comboBox and pushButton_Next_If_Doc_or_multi_Files
            self.ui.list_Sequence_titles_comboBox.activated.connect(self.show_selected_sequence)

            self.ui.pushButton_Next_If_Doc_or_multi_Files.clicked.connect(self.show_next_sequence)
            # ---- generate Create_ALL_in_One .docx file
            self.ui.pushButton_Create_ALL_in_One.clicked.connect(self.generate_Doc_Create_ALL_in_One)
        else:
            QMessageBox.warning(self, "No Selection", "No file selected.")

    def rename_file(self):
        # Open a dialog to select a file
        file_name, _ = QFileDialog.getOpenFileName(self, "Select File", "", "All Files (*);;Text Files (*.txt);;FASTA "
                                                                            "Files (*.fasta *.fas *.fa "
                                                                            "*.fna)")

        if file_name:
            # Get the original file extension
            original_extension = os.path.splitext(file_name)[1]

            # Get the new file name from the user
            new_name, ok = QtWidgets.QInputDialog.getText(self, 'File Rename',
                                                          'Enter new file name (without extension):')

            if ok and new_name:
                # Append the original file extension to the new name
                new_name_with_extension = f"{new_name}{original_extension}"

                # Create a new file path with the new name in the same directory
                new_file_path = os.path.join(os.path.dirname(file_name), new_name_with_extension)

                try:
                    # Rename the file
                    os.rename(file_name, new_file_path)
                    QMessageBox.information(self, "Success", "File renamed successfully!")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to rename file: {e}")
            else:
                QMessageBox.warning(self, "Cancelled", "File renaming cancelled.")
        else:
            QMessageBox.warning(self, "No Selection", "There is no File Name")

    # Message Error/ Information/Critical , I was using it before in the Previous Version of AL CODON APP
    # but Now, I AM not using it any more , I Found using"QMessageBox" directly way simpler :))
    def show_message(self, title, message, icon):
        msg = QMessageBox(self)
        msg.setWindowTitle(title)
        msg.setText(message)

        if icon == "information":
            msg.setIcon(QMessageBox.Information)
        elif icon == "critical":
            msg.setIcon(QMessageBox.Critical)
        elif icon == "warning":
            msg.setIcon(QMessageBox.Warning)

        # Apply custom styles
        msg.setStyleSheet("""
        QMessageBox {
            background-color: #2c3e50;  /* Dark background for a sleek look */
            font-size: 14px;
            border-radius: 10px;
            padding: 20px;  /* Increase padding for a larger box */
            min-width: 300px;  /* Set a minimum width */
            min-height: 150px;  /* Set a minimum height */
        }
        QMessageBox QLabel {
            color: #ecf0f1;  /* Light text for contrast */
            font-size: 14px;
            margin-bottom: 20px;  /* Add margin at the bottom of the text */
        }
        QMessageBox QPushButton {
            background-color: #2980b9;  /* Blue buttons */
            color: white;  /* White text on buttons */
            border: none;  /* No border */
            border-radius: 5px;  /* Rounded corners */
            padding: 10px 20px;  /* Increase padding for buttons */
            font-size: 14px;
            min-width: 100px;  /* Ensure buttons are a consistent size */
        }
        QMessageBox QPushButton:hover {
            background-color: #3498db;  /* Lighter blue on hover */
        }
        QMessageBox QPushButton:pressed {
            background-color: #1f618d;  /* Darker blue when pressed */
        }
        """)

        msg.exec_()

    def show_ANY_file(self, fname):

        # Read the file and handle the different return cases
        self.result_file = self.counter.read_ANY_file(fname)
        # Extract mandatory values first
        text_File = self.result_file[0]
        file_new_title = self.result_file[1]

        self.additional_text = self.result_file[2]

        self.titles = self.result_file[3]
        self.cleaned_sequences = self.result_file[4]
        doc_exit = self.result_file[5]
        nb_sequences = self.result_file[6]
        if nb_sequences > 1:
            QMessageBox.critical(self, "Error",
                                 "The selected File contains more than ONE Sequence please select "
                                 "'Select Multiple (.txt/.Fasta) Files' if you want to Process Multi "
                                 "Sequences ")
        else:
            # Populate combo box and store sequence data
            self.combo_titles = self.titles

            self.sequence_data = dict(zip(self.titles, self.cleaned_sequences))

            # Show the first sequence

            if doc_exit:

                self.ui.DNA_Sequnce_input.setPlainText(self.sequence_data[self.titles[0]])

                self.ui.selected_fileshow.setText(file_new_title)

                self.ui.textEdit_additional_text.setText(self.additional_text[0])

                self.ui.list_Sequence_titles_comboBox.clear()

                # Add a list_Sequence_titles
                self.ui.list_Sequence_titles_comboBox.addItems(self.titles)

                # Make list_Sequence_titles_comboBox and pushButton_Next_If_Doc_or_multi_Files
                self.ui.list_Sequence_titles_comboBox.activated.connect(self.show_selected_sequence)
                self.ui.pushButton_Next_If_Doc_or_multi_Files.clicked.connect(self.show_next_sequence)
                # ---- generate Create_ALL_in_One .docx file
                self.ui.pushButton_Create_ALL_in_One.clicked.connect(self.generate_Doc_Create_ALL_in_One)

            else:
                # Show the selected File FIRST Sewquence, Add 'Additional Infos' and 'Titles" if they exist
                self.ui.DNA_Sequnce_input.setPlainText(text_File)
                self.ui.selected_fileshow.setText(file_new_title)
                self.ui.textEdit_additional_text.setText(self.additional_text[0])

    #  Connect the 'pushButton_Next_If_Doc_or_multi_Files' and 'pushButton_Create_ALL_in_One'
    #  button's clicked signal to the slot , to Handle The ERROR messages if clicked when Sequence Input is empty
    def handle_button_click(self):
        if not self.ui.DNA_Sequnce_input.toPlainText().strip():
            QMessageBox.warning(self, "Input Check", "The Sequence Input is empty.\n"
                                                     "It seems that there is No File selected.")

    def show_selected_sequence(self):
        # delete text in content
        self.ui.selected_fileshow.setText('Loading NEXT Sequence')
        self.ui.DNA_Sequnce_output.clear()
        self.ui.textEdit_results.clear()
        self.ui.textEdit_results_TAIL.clear()
        title = self.ui.list_Sequence_titles_comboBox.currentText()

        if title in self.sequence_data:

            # Find the index of the title in the original titles list
            title_index = self.titles.index(title)

            # ----------------------------------------------------------------------------------------------------------
            # Bro I Need to find a better solution here !!!!!!!!!!!!!!
            # If self.sequence_data[next_title] is not empty, execute the code here
            if self.sequence_data[title]:
                # --------------------------------------------------

                if len(self.sequence_data[title]) <= 112000:
                    self.ui.DNA_Sequnce_input.setPlainText(self.sequence_data[title])
                    self.ui.textEdit_additional_text.setText(self.additional_text[title_index])
                    self.ui.selected_fileshow.setText(title)

                # --------------------------------------------------

                else:
                    # --------------------
                    # process the sequence:
                    #
                    # Prepares a task with the selected title, sequence data, and additional text.
                    # Uses a multiprocessing Pool to process the task asynchronously.
                    # Starts a QTimer to periodically check for results.
                    # Prepare task
                    task = (title, self.sequence_data, self.additional_text[title_index])

                    # Increase the number of processes
                    pool = Pool(processes=6)
                    pool.apply_async(process_data, (task,), callback=self.handle_result)
                    pool.close()
                    pool.join()

                    # Start the timer to check for results periodically
                    if not self.timer.isActive():
                        # Check every 100 ms
                        self.timer.start(100)

            else:
                # NEED to TEST THIS IF there is NOOOOO Sequence (EMPTY)
                # If self.sequence_data[next_title] is empty, execute the code here
                error_text = self.additional_text[title_index]

                text_input, additional_text = self.counter_docx.count_nucleotide_sequences_Docx_RECTIF(error_text)

                self.ui.DNA_Sequnce_input.setPlainText(text_input)
                self.ui.textEdit_additional_text.setText(additional_text)
                self.ui.selected_fileshow.setText(title)

            # ---------------------------------------------------------------------------------------------------------

    def show_next_sequence(self):
        # delete text in content
        self.ui.DNA_Sequnce_output.clear()
        self.ui.textEdit_results.clear()
        self.ui.textEdit_results_TAIL.clear()

        if self.current_index < len(self.combo_titles) - 1:
            self.current_index += 1
        else:
            self.current_index = 0

        next_title = self.combo_titles[self.current_index]

        # Bro I Need to find a better solution here !!!!!!!!!!!!!!
        # ------------------------------------------------------------
        self.ui.list_Sequence_titles_comboBox.setCurrentText(next_title)
        # If self.sequence_data[next_title] is not empty, execute the code here
        if self.sequence_data[next_title]:

            # --------------------------------------------------

            if len(self.sequence_data[next_title]) <= 112000:
                self.ui.DNA_Sequnce_input.setPlainText(self.sequence_data[next_title])
                self.ui.textEdit_additional_text.setText(self.additional_text[self.current_index])
                self.ui.selected_fileshow.setText(next_title)

            # --------------------------------------------------

            else:
                self.ui.selected_fileshow.setText('Loading NEXT Sequence')
                # --------------------
                # process the sequence:
                #
                # Prepares a task with the selected title, sequence data, and additional text.
                # Uses a multiprocessing Pool to process the task asynchronously.
                # Starts a QTimer to periodically check for results.
                # Prepare task
                task = (next_title, self.sequence_data, self.additional_text[self.current_index])

                # Increase the number of processes
                pool = Pool(processes=6)
                pool.apply_async(process_data, (task,), callback=self.handle_result)
                pool.close()
                pool.join()

                # Start the timer to check for results periodically
                if not self.timer.isActive():
                    # Check every 100 ms
                    self.timer.start(100)
        else:
            # If self.sequence_data[next_title] is empty, execute the code here
            error_text = self.additional_text[self.current_index]

            text_input, additional_text = self.counter_docx.count_nucleotide_sequences_Docx_RECTIF(error_text)

            self.ui.DNA_Sequnce_input.setPlainText(text_input)
            self.ui.textEdit_additional_text.setText(additional_text)
            self.ui.selected_fileshow.setText(next_title)
        # -----------------------------------------------------------------

    def Launch(self):
        # Define the strings to search for
        search_strings = ["ATTCA", "TTCAA", "TCAAG", "CAAGA", "AAGAT", "AGATG", "GATGA", "ATGAA", "TGAAT"]

        # Map each search string to a specific color
        color_mapping = {
            "ATTCA": "red",
            "TTCAA": "blue",
            "TCAAG": "green",
            "CAAGA": "purple",
            "AAGAT": "orange",
            "AGATG": "pink",
            "GATGA": "brown",
            "ATGAA": "teal",
            "TGAAT": "gray"
        }
        text_input = self.ui.DNA_Sequnce_input.toPlainText()

        results = self.counter.find_and_count_with_coloring(text_input, search_strings, color_mapping)
        results_TAIL = self.counter.find_and_count_with_coloring_TAIL(text_input)
        # Store results if needed elsewhere in the class
        # results[3] contains the styled/colored text, HTML-formatted strings with color styling,
        # I should use the setHtml() method instead of setPlainText()
        # to maintain the color formatting.
        self.ui.DNA_Sequnce_output.setHtml(results[3])

        # Call the Display_results method and set its return value to the textEdit_results/textEdit_results_TAIL

        results_summary = self.counter.Display_results(results)
        results_summary_TAIL = self.counter.Display_results_TAIL(results_TAIL)

        self.ui.textEdit_results.setPlainText(results_summary)
        self.ui.textEdit_results_TAIL.setPlainText(results_summary_TAIL)
        return results, results_TAIL

    def handle_launch(self):
        # Check if the DNA_Sequnce_input is empty(self, "Input Check", "The Sequence Input is empty.\n"
        #                                                      "It seems that there is No File selected.")
        if not self.ui.DNA_Sequnce_input.toPlainText().strip():
            QMessageBox.warning(self, "Input Check", "The Sequence Input is empty.\n"
                                                     "It seems that there is No File selected.")
        else:
            self.results, self.results_TAIL = self.Launch()
            # Now you can use 'results' here or store it in the class for other uses
            self.stored_results = self.results
            self.clicker_list_pentamers
            self.clicker_list_pentamers_TAIL

    # -------------------------Handle Multi Processing and Thread ------------------------------------------
    # ------

    # handle_result:
    #
    # Puts the processed result into the results queue.
    def handle_result(self, result):

        self.results_queue.put(result)

    # check_results:
    #
    # Periodically checks the results queue for new results.
    # Updates the GUI with the results if available.
    def check_results(self):
        while not self.results_queue.empty():
            result = self.results_queue.get()
            self.update_gui(*result)
        if self.results_queue.empty():
            self.timer.stop()

    # update_gui:
    #
    # Updates the GUI elements with the processed data.
    def update_gui(self, title, processed_sequence, additional_text):

        self.ui.DNA_Sequnce_input.setPlainText(processed_sequence)
        self.ui.textEdit_additional_text.setText(additional_text)
        self.ui.selected_fileshow.setText(title)

    # closeEvent:
    #
    # Stops the result listener thread when the window is closed.
    def closeEvent(self, event):
        self.result_listener.stop()
        self.result_listener.wait()

        super().closeEvent(event)

    # ------------------------------------------------------------------------------------------------------


if __name__ == "__main__":
    freeze_support()
    app = QApplication(sys.argv)

    # Create and show the splash screen
    splash = SplashScreen()
    sys.exit(app.exec_())
